"""
Generate a formatted Word document (report.docx) with synthesis memo,
embedded charts, and notable quotes appendix.
"""

import json
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUTPUT_DIR = Path(__file__).parent / "output"
CHARTS_DIR = OUTPUT_DIR / "charts"


def load_data():
    """Load analyzed comments and synthesis memo."""
    analyzed_file = OUTPUT_DIR / "analyzed_comments.json"
    memo_file = OUTPUT_DIR / "synthesis_memo.txt"

    analyzed = []
    memo = ""

    if analyzed_file.exists():
        with open(analyzed_file, "r", encoding="utf-8") as f:
            analyzed = json.load(f)

    if memo_file.exists():
        with open(memo_file, "r", encoding="utf-8") as f:
            memo = f.read()

    return analyzed, memo


def generate_charts(analyzed: list[dict]):
    """Generate chart PNGs from analyzed data."""
    CHARTS_DIR.mkdir(parents=True, exist_ok=True)

    plt.style.use("dark_background")
    colors_ngne = "#3b82f6"
    colors_tsha = "#10b981"

    # 1. Sentiment bar chart
    fig, ax = plt.subplots(figsize=(8, 4))
    categories = ["Positive", "Negative", "Neutral"]
    ngne_vals = [
        sum(1 for c in analyzed if c.get("sentiment_toward_NGNE") == cat.lower())
        for cat in categories
    ]
    tsha_vals = [
        sum(1 for c in analyzed if c.get("sentiment_toward_TSHA") == cat.lower())
        for cat in categories
    ]
    x = range(len(categories))
    width = 0.35
    ax.bar([i - width / 2 for i in x], ngne_vals, width, label="NGNE (Neurogene)", color=colors_ngne)
    ax.bar([i + width / 2 for i in x], tsha_vals, width, label="TSHA (Taysha)", color=colors_tsha)
    ax.set_ylabel("Comment Count")
    ax.set_title("Sentiment Distribution: NGNE vs TSHA")
    ax.set_xticks(list(x))
    ax.set_xticklabels(categories)
    ax.legend()
    plt.tight_layout()
    fig.savefig(CHARTS_DIR / "sentiment_bar.png", dpi=150, bbox_inches="tight")
    plt.close()

    # 2. Implied lean pie chart
    fig, ax = plt.subplots(figsize=(6, 6))
    lean_counts = {}
    for c in analyzed:
        lean = c.get("implied_lean", "unclear")
        lean_counts[lean] = lean_counts.get(lean, 0) + 1
    labels = list(lean_counts.keys())
    sizes = list(lean_counts.values())
    color_map = {"NGNE": colors_ngne, "TSHA": colors_tsha, "neither": "#64748b", "unclear": "#f59e0b"}
    pie_colors = [color_map.get(l, "#94a3b8") for l in labels]
    ax.pie(sizes, labels=labels, autopct="%1.1f%%", colors=pie_colors, textprops={"color": "white"})
    ax.set_title("Implied Therapy Lean Distribution")
    plt.tight_layout()
    fig.savefig(CHARTS_DIR / "lean_pie.png", dpi=150, bbox_inches="tight")
    plt.close()

    # 3. Theme frequency horizontal bar
    theme_counts = {}
    for c in analyzed:
        for theme in c.get("themes", []):
            theme_counts[theme] = theme_counts.get(theme, 0) + 1
    sorted_themes = sorted(theme_counts.items(), key=lambda x: x[1], reverse=True)[:15]

    if sorted_themes:
        fig, ax = plt.subplots(figsize=(8, 6))
        themes = [t[0].replace("_", " ") for t in sorted_themes]
        counts = [t[1] for t in sorted_themes]
        ax.barh(themes[::-1], counts[::-1], color=colors_ngne)
        ax.set_xlabel("Frequency")
        ax.set_title("Top Discussion Themes in RTT Community")
        plt.tight_layout()
        fig.savefig(CHARTS_DIR / "themes_bar.png", dpi=150, bbox_inches="tight")
        plt.close()

    print(f"Generated chart PNGs in {CHARTS_DIR}")


def build_docx(analyzed: list[dict], memo: str):
    """Build the Word document."""
    doc = Document()

    # -- Styles --
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # -- Cover Page --
    for _ in range(6):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("RTT Community Sentiment Analysis")
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Alternative Data Report")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run(
        "Facebook Community Discourse Analysis\n"
        "Neurogene (NGNE / NGN-401) vs. Taysha (TSHA / TSHA-102)\n"
        "Gene Therapy for Rett Syndrome"
    )
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    for _ in range(4):
        doc.add_paragraph("")

    stats_para = doc.add_paragraph()
    stats_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    total = len(analyzed)
    notable_count = sum(1 for c in analyzed if c.get("notable_quote"))
    run = stats_para.add_run(f"Based on analysis of {total} community comments | {notable_count} notable quotes identified")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    doc.add_page_break()

    # -- Table of Contents placeholder --
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Overall Community Sentiment",
        "2. Therapy Preference Signals",
        "3. Clinically Unaddressed Concerns",
        "4. The ICV Question",
        "5. Reaction to the NGNE Patient Death",
        "6. Implications for Commercial Uptake",
        "Appendix: Notable Quotes",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # -- Embedded Charts --
    doc.add_heading("Key Charts", level=1)

    chart_files = [
        ("sentiment_bar.png", "Sentiment Distribution: NGNE vs TSHA"),
        ("lean_pie.png", "Implied Therapy Lean"),
        ("themes_bar.png", "Top Discussion Themes"),
    ]

    for fname, caption in chart_files:
        chart_path = CHARTS_DIR / fname
        if chart_path.exists():
            doc.add_paragraph(caption).bold = True
            doc.add_picture(str(chart_path), width=Inches(5.5))
            doc.add_paragraph("")

    doc.add_page_break()

    # -- Synthesis Memo --
    doc.add_heading("Synthesis Research Memo", level=1)

    if memo:
        # Parse memo sections
        lines = memo.split("\n")
        for line in lines:
            line = line.rstrip()
            if not line:
                doc.add_paragraph("")
            elif line.startswith("**") and line.endswith("**"):
                doc.add_heading(line.strip("* "), level=2)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("# "):
                doc.add_heading(line[2:], level=2)
            elif line.startswith("> "):
                p = doc.add_paragraph()
                p.style = "Quote" if "Quote" in [s.name for s in doc.styles] else "Normal"
                run = p.add_run(line[2:])
                run.italic = True
            elif line.startswith("- "):
                doc.add_paragraph(line[2:], style="List Bullet")
            else:
                # Handle bold text within lines
                doc.add_paragraph(line)
    else:
        doc.add_paragraph("(Synthesis memo not yet generated. Run analyzer.py first.)")

    doc.add_page_break()

    # -- Appendix: Notable Quotes --
    doc.add_heading("Appendix: Top Notable Quotes", level=1)

    notable = [c for c in analyzed if c.get("notable_quote")][:30]

    if notable:
        table = doc.add_table(rows=1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        # Header
        headers = ["Author", "Quote", "Themes", "Implied Lean"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)

        # Data rows
        for q in notable:
            row = table.add_row()
            row.cells[0].text = q.get("author_anon", "Unknown")
            quote_text = q.get("verbatim_quote") or q.get("text", "")
            row.cells[1].text = quote_text[:300]
            row.cells[2].text = ", ".join(q.get("themes", []))
            row.cells[3].text = q.get("implied_lean", "unclear")

            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)

        # Set column widths
        for row in table.rows:
            row.cells[0].width = Cm(2.5)
            row.cells[1].width = Cm(10)
            row.cells[2].width = Cm(3.5)
            row.cells[3].width = Cm(2)
    else:
        doc.add_paragraph("No notable quotes identified.")

    # Save
    report_path = OUTPUT_DIR / "report.docx"
    doc.save(str(report_path))
    print(f"Wrote {report_path}")


def generate_report():
    """Main entry point."""
    analyzed, memo = load_data()
    if not analyzed:
        print("No analyzed data found. Run the pipeline first.")
        return

    generate_charts(analyzed)
    build_docx(analyzed, memo)


if __name__ == "__main__":
    generate_report()
